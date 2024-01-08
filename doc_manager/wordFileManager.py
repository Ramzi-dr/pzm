import os
import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm
from docx.shared import Pt
from docx.shared import RGBColor
from doc_manager.monthOrYearTableManager import MonthOrYearTableManager
from doc_manager.weekTableManager import WeekTableManager
from doc_manager.yearTableManager import YearTableManager
from emailManager import send_email


class WordFileManager:
    def __init__(
        self,
        door_Fullname,
        door_name,
        door_number,
        door_house,
        door_floor,
        current_year,
        current_month,
        current_week_number,
    ):
        self.door_FULLname = door_Fullname
        self.door_number = door_number
        self.door_name = door_name
        self.door_house = door_house
        self.door_floor = door_floor
        self.year = current_year
        self.month = current_month
        self.week_number = current_week_number
        self.headWeekTitle_text = f"Wochen-Auswertung: {current_week_number}"
        self.headMonthTitle_text = f"Monats-Auswertung: {current_month}"
        self.headYearTitle_text = f"Jahres-Auswertung: {current_year}"

        self.list_headTitle_text = [
            self.headWeekTitle_text,
            self.headMonthTitle_text,
            self.headYearTitle_text,
        ]
        self.week_title_text = (
            f"Auswertung {door_name} {door_house} {current_week_number}"
        )
        self.month_title_text = (
            f"Auswertung {door_name} {door_house} des Monats {current_month}"
        )
        self.year_title_text = (
            f"Auswertung {door_name} {door_house} des Jahres {current_year}"
        )

        self.list_title_text = [
            self.week_title_text,
            self.month_title_text,
            self.year_title_text,
        ]
        self.floor_text = f"Stockwerk:   {door_floor}"
        self.number_text = f"Tür-Nr:          {door_number}"

        self.text_toChange = ["{main_title}", "{door_name}"]

        self.folder_path = "doors_wordFiles"
        self.weekFilename = (
            f"{door_Fullname} {current_year} {current_week_number} .docx"
        )
        self.monthFilename = f"{door_Fullname} {current_year} {current_month} .docx"
        self.yearFilename = f"{door_Fullname} {current_year} .docx"
        self.list_filename = [
            self.weekFilename,
            self.monthFilename,
            self.yearFilename,
        ]
        self.weekFile_path = os.path.join(self.folder_path, self.weekFilename)
        self.monthFile_path = os.path.join(self.folder_path, self.monthFilename)
        self.yearFile_path = os.path.join(self.folder_path, self.yearFilename)

        self.list_filePath = [
            self.weekFile_path,
            self.monthFile_path,
            self.yearFile_path,
        ]
        self.left_alignment = WD_ALIGN_PARAGRAPH.RIGHT
        self.right_alignment = WD_ALIGN_PARAGRAPH.LEFT
        self.center_alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.head_text_color = RGBColor(0, 0, 0)
        self.title_text_color = RGBColor(215, 133, 131)

    def check_fileExist(self, filename):
        try:
            if os.path.exists(self.folder_path):
                file_exist = os.path.exists(os.path.join(self.folder_path, filename))
                return file_exist
        except Exception as e:
            print(f"there is error : {e}")

    def createOrUpdate_File(self, close_time, open_time, percent_close, percent_open):
        week_doc = 0
        month_doc = 1
        year_doc = 2
        for i in range(len(self.list_filename)):
            if self.check_fileExist(filename=self.list_filename[i]):
                doc = docx.Document(self.list_filePath[i])
                if i == week_doc:
                    weekTable = WeekTableManager(
                        doc,
                        close_time=close_time[i],
                        open_time=open_time[i],
                        percent_close=percent_close[i],
                        percent_open=percent_open[i],
                    )
                    weekTable.createOrUpdate_table()

                if i == month_doc:
                    monthTable = MonthOrYearTableManager(
                        doc=doc,
                        close_time=close_time[2],
                        open_time=open_time[2],
                        percent_close=percent_close[2],
                        percent_open=percent_open[2],
                        table_info_text=self.month,
                    )
                    monthTable.createOrUpdate_table()

                if i == year_doc:
                    yearTable = YearTableManager(
                        doc=doc,
                        the_month=self.month,
                        the_year=self.year,
                        close_time=close_time[2],
                        year_close_time=close_time[3],
                        open_time=open_time[2],
                        year_open_time=open_time[3],
                        percent_close=percent_close[2],
                        year_percent_close=percent_close[3],
                        percent_open=percent_open[2],
                        year_percent_open=percent_open[3],
                    )
                    yearTable.createOrUpdate_table()
                try:
                    doc.save(self.list_filePath[i])
                except PermissionError:
                    send_email(subject='Error',message=f'error in Pzm Event Server in createOrUpdate_File() at wordFIleManager.py\n {e}  ')
                    pass
                except Exception as e:
                    send_email(subject='Error',message=f'error in Pzm Event Server in createOrUpdate_File() at wordFIleManager.py\n {e}  ')
                    pass

            else:
                doc = docx.Document()
                headTitle = doc.add_paragraph(self.list_headTitle_text[i])
                headTitle.alignment = self.right_alignment
                headTitle.runs[0].font.size = Pt(10)
                doc.add_picture("/home/adminbst/Documents/pzmCode/doors_wordFiles/pzmPicture.png")
                pic_1 = doc.paragraphs[-1]
                pic_1.alignment = self.left_alignment
                head1 = doc.add_paragraph(self.list_title_text[i])
                head2 = doc.add_paragraph(self.floor_text)
                head3 = doc.add_paragraph(self.number_text)
                head1.style.font.color.rgb = self.head_text_color
                head1.style.font.size = Pt(16)
                head1.style.font.name = "Arial"
                head1.paragraph_format.space_before = Pt(50)
                head1.paragraph_format.space_after = Pt(0.1)
                head2.style.font.color.rgb = self.head_text_color
                head2.style.font.size = Pt(16)
                head1.style.font.name = "Arial"
                head2.paragraph_format.space_before = Pt(0)
                head2.paragraph_format.space_after = Pt(0.1)
                head3.runs[0].style.font.color.rgb = self.head_text_color
                head3.style.font.size = Pt(16)
                head1.style.font.name = "Arial"
                head3.paragraph_format.space_before = Pt(0)
                head3.paragraph_format.space_after = Pt(10)

                par1 = doc.add_paragraph(
                    "Wenn die Hauptzugangstür verriegelt ist, befinden sich die Patienten / Bewohner in einer geschlossenen Abteiltung."
                    " Es erfolgt eine automatische Erfassung über den Verriegelungszustand."
                )
                par1.runs[0].font.size = Pt(10)
                par1.runs[0].font.bold = False
                par2 = doc.add_paragraph("Mögliche Betriebszustände:")

                par2.paragraph_format.space_before = Pt(20)
                par2.runs[0].font.size = Pt(10)
                line1 = doc.add_paragraph(
                    "Zugang geschlossen (rot markiert)", style="List Bullet"
                )
                line1.runs[0].font.size = Pt(10)
                line1.paragraph_format.space_after = Pt(0)
                line2 = doc.add_paragraph(
                    "    Die Tür ist verriegelt und kann nur mit Badge / Schlüssel geöffnet werden"
                )
                line2.runs[0].font.size = Pt(10)
                line2.paragraph_format.space_before = Pt(0)
                line2.paragraph_format.space_after = Pt(0)
                line3 = doc.add_paragraph(
                    "Zugang geöffnet (grün markiert)", style="List Bullet"
                )
                line3.runs[0].font.size = Pt(10)
                line3.paragraph_format.space_before = Pt(0)
                line3.paragraph_format.space_after = Pt(0)
                line4 = doc.add_paragraph(
                    "    Die Tür ist geöffnet und kann beidseitig ohne Badge / Schlüssel geöffnet werden"
                )
                line4.runs[0].font.size = Pt(10)
                line4.paragraph_format.space_before = Pt(0)
                week_title1 = f"Auswert-Zeitraum:  {self.week_number} - {self.year} (08:00 - 18:30)"
                month_title1 = (
                    f"Auswert-Zeitraum:  {self.month} - {self.year} (08:00 - 18:30)"
                )
                year_title1 = f"Auswert-Zeitraum:  {self.year} (08:00 - 18:30)"

                title1_list = [
                    week_title1,
                    month_title1,
                    year_title1,
                ]
                title1 = doc.add_paragraph(title1_list[i])
                title1.runs[0].font.size = Pt(13)
                title1.runs[0].font.bold = True
                title1.alignment = self.center_alignment
                for run in title1.runs:
                    run.font.color.rgb = self.title_text_color
                sections = doc.sections
                for section in sections:
                    section.top_margin = Cm(1)
                    section.bottom_margin = Cm(1)
                    section.left_margin = Cm(1)
                    section.right_margin = Cm(1)

                if i == week_doc:
                    weekTable = WeekTableManager(
                        doc=doc,
                        close_time=close_time[i],
                        open_time=open_time[i],
                        percent_close=percent_close[i],
                        percent_open=percent_open[i],
                    )
                    weekTable.createOrUpdate_table()
                if i == month_doc:
                    monthTable = MonthOrYearTableManager(
                        doc=doc,
                        close_time=close_time[2],
                        open_time=open_time[2],
                        percent_close=percent_close[2],
                        percent_open=percent_open[2],
                        table_info_text=f"{self.month} - {self.year}",
                    )
                    monthTable.createOrUpdate_table()

                elif i == year_doc:
                    yearTable = YearTableManager(
                        doc=doc,
                        the_month=self.month,
                        the_year=self.year,
                        close_time=close_time[2],
                        year_close_time=close_time[3],
                        open_time=open_time[2],
                        year_open_time=open_time[3],
                        percent_close=percent_close[2],
                        year_percent_close=percent_close[3],
                        percent_open=percent_open[2],
                        year_percent_open=percent_open[3],
                    )
                    yearTable.createOrUpdate_table()

                doc.save(self.list_filePath[i])
