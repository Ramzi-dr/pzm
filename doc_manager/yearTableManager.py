from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from doc_manager.monthOrYearTableManager import MonthOrYearTableManager
from payloadCollection import PayloadCollection
from .chartCreator import ChartCreator


class YearTableManager:
    def __init__(
        self,
        doc,
        the_month,
        the_year,
        close_time,
        year_close_time,
        open_time,
        year_open_time,
        percent_close,
        year_percent_close,
        percent_open,
        year_percent_open,
    ):
        self.doc = doc
        self.the_month = the_month
        self.the_year = the_year
        self.close_time = close_time
        self.year_close_time = year_close_time
        self.open_time = open_time
        self.year_open_time = year_open_time
        self.percent_close = percent_close
        self.year_percent_close = year_percent_close
        self.percent_open = percent_open
        self.year_percent_open = year_percent_open
        self.first_6_monthsOfYear = PayloadCollection.german_months[:6]
        self.second_6_monthsOfYear = PayloadCollection.german_months[6:]
        self.left_alignment = WD_ALIGN_PARAGRAPH.RIGHT
        self.right_alignment = WD_ALIGN_PARAGRAPH.LEFT
        self.center_alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    def update_year_table(self, year_table):
        chart_instance = ChartCreator(
            percent_open=self.year_percent_open, percent_close=self.year_percent_close
        )
        chart = chart_instance.generate_chart()
        year_table.cell(1, 1).text = self.year_close_time
        year_table.cell(2, 1).text = self.year_open_time
        year_table.cell(3, 1).text = self.year_percent_close
        year_table.cell(4, 1).text = self.year_percent_open
        chart_cell = year_table.cell(5, 1)
        for para in chart_cell.paragraphs:
            for run in para.runs:
                run.clear()

        if not chart_cell.paragraphs:  # Check if the cell is empty
            chart_cell.paragraphs.append("")  # Add an empty paragraph

        chart_cell.paragraphs[0].paragraph_format.alignment = WD_ALIGN_VERTICAL.CENTER
        chart_cell.paragraphs[0].add_run().add_picture(chart, width=Cm(5))

    def year_tableGenerator(self):
        year_table = MonthOrYearTableManager(
            doc=self.doc,
            close_time=self.year_close_time,
            open_time=self.year_open_time,
            percent_close=self.year_percent_close,
            percent_open=self.year_percent_open,
            table_info_text=self.the_year,
        )
        yearTable = year_table.createOrUpdate_table()
        return yearTable

    def update_halfYear_table(self, halfYear_table, cell_toUpdate):
        chart_instance = ChartCreator(
            percent_open=self.percent_open,
            percent_close=self.percent_close,
        )
        chart = chart_instance.generate_chart()
        halfYear_table.cell(1, cell_toUpdate).text = self.close_time
        halfYear_table.cell(2, cell_toUpdate).text = self.open_time
        halfYear_table.cell(3, cell_toUpdate).text = self.percent_close
        halfYear_table.cell(4, cell_toUpdate).text = self.percent_open
        chart_cell = halfYear_table.cell(5, cell_toUpdate)
        for para in chart_cell.paragraphs:
            for run in para.runs:
                run.clear()

        if not chart_cell.paragraphs:  # Check if the cell is empty
            chart_cell.paragraphs.append("")  # Add an empty paragraph

        chart_cell.paragraphs[0].paragraph_format.alignment = WD_ALIGN_VERTICAL.CENTER
        chart_cell.paragraphs[0].add_run().add_picture(chart, width=Cm(2))

    def halfYear_tableGenerator(self, list_of_month):
        chart_instance = ChartCreator(
            percent_open=self.percent_open, percent_close=self.percent_close
        )
        chart = chart_instance.generate_chart()
        half_year = self.doc.add_table(rows=6, cols=7, style="Table Grid")
        row_months = half_year.rows[0].cells
        row_closeTime = half_year.rows[1].cells
        row_closeTime[0].text = "Zeit geschlossen"
        row_openTime = half_year.rows[2].cells
        row_openTime[0].text = "Zeit geöffnet"
        row_percentClose = half_year.rows[3].cells
        row_percentClose[0].text = "Prozentsatz\u00A0geschlossen"
        row_percentOpen = half_year.rows[4].cells
        row_percentOpen[0].text = "Prozentsatz\u00A0geöffnet"
        row_charts = half_year.rows[5].cells
        row_charts[0].text = "Diagramm"

        for i in range(len(row_months)):
            if i == 0:
                row_months[i].text = PayloadCollection.weekTable_first_row[i]

            else:
                row_months[i].text = list_of_month[i - 1]
            if self.the_month in list_of_month:
                if row_months[i].text == self.the_month:
                    row_closeTime[i].text = self.close_time
                    row_openTime[i].text = self.open_time
                    row_percentOpen[i].text = self.percent_open
                    row_percentClose[i].text = self.percent_close
                    chart_cell = row_charts[i]
                    chart_cell.paragraphs[
                        0
                    ].paragraph_format.alignment = WD_ALIGN_VERTICAL.CENTER

                    chart_cell.paragraphs[0].add_run().add_picture(chart, width=Cm(2))
        for i, row in enumerate(half_year.rows):
            if i == len(half_year.rows) - 1:
                # Set a different height for the last row
                row.height = Cm(2.0)  # Adjust the height as needed
                for cell in row.cells:
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            else:
                row.height = Cm(1)  # Set the height for all other rows

        for row in half_year.rows:
            for cell in row.cells:
                cell.vertical_alignment = self.vertical_alignment

                paragraphs = cell.paragraphs
                for paragraph in paragraphs:
                    if paragraph.text not in [
                        "Totalisierung",
                        "Zeit geschlossen",
                        "Zeit geöffnet",
                        "Prozentsatz\u00A0geschlossen",
                        "Prozentsatz\u00A0geöffnet",
                        "Diagramm",
                    ]:
                        paragraph.alignment = self.center_alignment

                    for run in paragraph.runs:
                        run.font.size = Pt(10)
        # Access the first row of the table
        first_row = half_year.rows[0]

        # Loop through the cells in the first row
        for cell in first_row.cells:
            # Create a shading object with the desired background color (e.g., light blue)
            shading = cell._tc.get_or_add_tcPr()
            # CREATE SHADING OBJECT
            shade_obj = OxmlElement("w:shd")
            # SET THE SHADING OBJECT
            shade_obj.set(qn("w:fill"), "808080")
            # APPEND THE PROPERTIES TO THE TABLE CELL PROPERTIES
            shading.append(shade_obj)
            # Access the paragraph inside the cell
            paragraph = cell.paragraphs[0]

            # Access the existing text and make it bold
            for run in paragraph.runs:
                run.bold = True
        return half_year

    def createOrUpdate_table(self):
        if len(self.doc.tables) > 0:
            for table in self.doc.tables:
                if len(table.columns) == 2:
                    year_table = table
                    self.update_year_table(year_table)
                elif len(table.columns) == 7:
                    for i in range(len(table.columns)):
                        if table.cell(0, i).text == self.the_month:
                            halfYear_table = table
                            self.update_halfYear_table(halfYear_table, i)

        else:
            table_of_theYear = self.year_tableGenerator()
            jump_to_newPage = self.doc.add_page_break()
            first_half = self.halfYear_tableGenerator(self.first_6_monthsOfYear)
            empty_space = self.doc.add_paragraph("")
            second_half = self.halfYear_tableGenerator(self.second_6_monthsOfYear)
            return (
                table_of_theYear,
                jump_to_newPage,
                first_half,
                empty_space,
                second_half,
            )
