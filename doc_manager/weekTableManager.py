from docx.shared import Pt
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from dateCalculator import get_dates_from_week, swiss_date
from payloadCollection import PayloadCollection
from .chartCreator import ChartCreator


class WeekTableManager:
    def __init__(self, doc, close_time, open_time, percent_close, percent_open):
        self.doc = doc
        self.the_date = swiss_date()
        self.formatted_date = self.the_date[:-4] + self.the_date[-2:]
        self.close_time = close_time
        self.open_time = open_time
        self.percent_close = percent_close
        self.percent_open = percent_open
        self.date_list, self.week_number = get_dates_from_week()
        self.left_alignment = WD_ALIGN_PARAGRAPH.RIGHT
        self.right_alignment = WD_ALIGN_PARAGRAPH.LEFT
        self.center_alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        self.no_break_space = "\u00A0"

    def createOrUpdate_table(self):
        chart_instance = ChartCreator(
            percent_open=self.percent_open,
            percent_close=self.percent_close
        )
        chart = chart_instance.generate_chart()

        if len(self.doc.tables) > 0:
            # Check if a table with the specified number of rows and columns already exists
            for table in self.doc.tables:
                if len(table.rows) == 6 and len(table.columns) == 8:
                    part2 = ""
                    for i in range(len(table.columns)):
                        text_parts = table.cell(0, i).text.split("\n")
                        if len(text_parts) == 2:
                            part1 = text_parts[0]
                            part2 = text_parts[1]
                        if part2 == self.formatted_date:
                            table.cell(1, i).text = self.close_time
                            table.cell(2, i).text = self.open_time
                            table.cell(3, i).text = self.percent_close
                            table.cell(4, i).text = self.percent_open
                            chart_cell = table.cell(5, i)
                            for para in chart_cell.paragraphs:
                                for run in para.runs:
                                    run.clear()

                            if not chart_cell.paragraphs:  # Check if the cell is empty
                                chart_cell.paragraphs.append(
                                    ""
                                )  # Add an empty paragraph

                            chart_cell.paragraphs[
                                0
                            ].paragraph_format.alignment = WD_ALIGN_VERTICAL.CENTER
                            chart_cell.paragraphs[0].add_run().add_picture(
                                chart, width=Cm(2)
                            )

        else:
            table = self.doc.add_table(rows=6, cols=8, style="Table Grid")
            row_days = table.rows[0].cells
            row_closeTime = table.rows[1].cells
            row_closeTime[0].text = "Zeit geschlossen"
            row_openTime = table.rows[2].cells
            row_openTime[0].text = "Zeit geöffnet"
            row_percentClose = table.rows[3].cells
            row_percentClose[0].text = "Prozentsatz\u00A0geschlossen"
            row_percentOpen = table.rows[4].cells
            row_percentOpen[0].text = "Prozentsatz\u00A0geöffnet"
            row_charts = table.rows[5].cells

            row_charts[0].text = "Diagramm"

            # Adding heading in the 1st row_days of the table
            for i in range(len(row_days)):
                if i == 0:
                    row_days[i].text = PayloadCollection.weekTable_first_row[i]

                else:
                    row_days[
                        i
                    ].text = f"{PayloadCollection.weekTable_first_row[i]}\n{self.date_list[i - 1]}"
                    if self.date_list[i - 1] == self.formatted_date:
                        row_openTime[i].text = self.open_time
                        row_percentOpen[i].text = self.percent_open
                        row_closeTime[i].text = self.close_time
                        row_percentClose[i].text = self.percent_close
                        chart_cell = row_charts[i]
                        chart_cell.paragraphs[
                            0
                        ].paragraph_format.alignment = WD_ALIGN_VERTICAL.CENTER

                        chart_cell.paragraphs[0].add_run().add_picture(
                            chart, width=Cm(2)
                        )

        for i, row in enumerate(table.rows):
            if i == len(table.rows) - 1:
                # Set a different height for the last row
                row.height = Cm(2.0)  # Adjust the height as needed
                for cell in row.cells:
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            else:
                row.height = Cm(1)  # Set the height for all other rows

        for row in table.rows:
            for cell in row.cells:
                cell.vertical_alignment = self.vertical_alignment

                paragraphs = cell.paragraphs
                for paragraph in paragraphs:
                    if paragraph.text not in [
                        "Totalisierung",
                        "Zeit geschlossen",
                        "Zeit geöffnet",
                        "Prozentsatz\u00A0geschlossen",
                        "Prozentsatz\u00A0geöffnet",
                        "Diagramm",
                    ]:
                        paragraph.alignment = self.center_alignment

                    for run in paragraph.runs:
                        run.font.size = Pt(10)
                        
        # Access the first row of the table
        first_row = table.rows[0]

        # Loop through the cells in the first row
        for cell in first_row.cells:
            # Create a shading object with the desired background color (e.g., light blue)
            shading = cell._tc.get_or_add_tcPr()
            # CREATE SHADING OBJECT
            shade_obj = OxmlElement("w:shd")
            # SET THE SHADING OBJECT
            shade_obj.set(qn("w:fill"), "808080")
            # APPEND THE PROPERTIES TO THE TABLE CELL PROPERTIES
            shading.append(shade_obj)
                # Access the paragraph inside the cell
            paragraph = cell.paragraphs[0]

            # Access the existing text and make it bold
            for run in paragraph.runs:
                run.bold = True



        return table
