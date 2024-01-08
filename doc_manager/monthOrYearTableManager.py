from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from .chartCreator import ChartCreator


class MonthOrYearTableManager:
    def __init__(
        self, doc, close_time, open_time, percent_close, percent_open, table_info_text
    ):
        self.doc = doc
        self.close_time = close_time
        self.open_time = open_time
        self.percent_close = percent_close
        self.percent_open = percent_open
        self.table_info_text = table_info_text
        self.left_alignment = WD_ALIGN_PARAGRAPH.RIGHT
        self.right_alignment = WD_ALIGN_PARAGRAPH.LEFT
        self.center_alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    def createOrUpdate_table(self):
        chart_instance = ChartCreator(
            percent_open=self.percent_open,
            percent_close=self.percent_close
        )
        chart = chart_instance.generate_chart()

        if len(self.doc.tables) > 0:
            # Check if a table with the specified number of rows and columns already exists
            for table in self.doc.tables:
                if len(table.rows) == 6 and len(table.columns) == 2:
                    table.cell(1, 1).text = self.close_time
                    table.cell(2, 1).text = self.open_time
                    table.cell(3, 1).text = self.percent_close
                    table.cell(4, 1).text = self.percent_open
                    chart_cell = table.cell(5, 1)
                    for para in chart_cell.paragraphs:
                        for run in para.runs:
                            run.clear()

                    if not chart_cell.paragraphs:  # Check if the cell is empty
                        chart_cell.paragraphs.append("")  # Add an empty paragraph

                    chart_cell.paragraphs[
                        0
                    ].paragraph_format.alignment = WD_ALIGN_VERTICAL.CENTER
                    chart_cell.paragraphs[0].add_run().add_picture(chart, width=Cm(5))

        else:
            table = self.doc.add_table(rows=6, cols=2, style="Table Grid")
            row_tableInfo = table.rows[0].cells
            row_tableInfo[0].text = "Totalisierung"
            row_tableInfo[1].text = f"{self.table_info_text}"
            row_closeTime = table.rows[1].cells
            row_closeTime[0].text = "Zeit geschlossen"
            row_closeTime[1].text = self.close_time
            row_openTime = table.rows[2].cells
            row_openTime[0].text = "Zeit geöffnet"
            row_openTime[1].text = self.open_time
            row_percentClose = table.rows[3].cells
            row_percentClose[0].text = "Prozentsatz\u00A0geschlossen"
            row_percentClose[1].text = self.percent_close
            row_percentOpen = table.rows[4].cells
            row_percentOpen[0].text = "Prozentsatz\u00A0geöffnet"
            row_percentOpen[1].text = self.percent_open
            row_charts = table.rows[5].cells
            row_charts[0].text = "Diagramm"
            chart_cell = row_charts[1]
            chart_cell.paragraphs[
                0
            ].paragraph_format.alignment = WD_ALIGN_VERTICAL.CENTER
            chart_cell.paragraphs[0].add_run().add_picture(chart, width=Cm(5))

        for i, row in enumerate(table.rows):
            if i == len(table.rows) - 1:
                # Set a different height for the last row
                row.height = Cm(2.0)  # Adjust the height as needed
                for cell in row.cells:
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            else:
                row.height = Cm(1)  # Set the height for all other rows

        for row in table.rows:
            for cell in row.cells:
                cell.vertical_alignment = self.vertical_alignment

                paragraphs = cell.paragraphs
                for paragraph in paragraphs:
                    if paragraph.text not in [
                        "Totalisierung",
                        "Zeit geschlossen",
                        "Zeit geöffnet",
                        "Prozentsatz\u00A0geschlossen",
                        "Prozentsatz\u00A0geöffnet",
                        "Diagramm",
                    ]:
                        paragraph.alignment = self.center_alignment

                    for run in paragraph.runs:
                        run.font.size = Pt(10)

        # Access the first row of the table
        first_row = table.rows[0]

        # Loop through the cells in the first row
        for cell in first_row.cells:
            # Create a shading object with the desired background color (e.g., light blue)
            shading = cell._tc.get_or_add_tcPr()
            # CREATE SHADING OBJECT
            shade_obj = OxmlElement("w:shd")
            # SET THE SHADING OBJECT
            shade_obj.set(qn("w:fill"), "808080")
            # APPEND THE PROPERTIES TO THE TABLE CELL PROPERTIES
            shading.append(shade_obj)
            # Access the paragraph inside the cell
            paragraph = cell.paragraphs[0]

            # Access the existing text and make it bold
            for run in paragraph.runs:
                run.bold = True

        return table
